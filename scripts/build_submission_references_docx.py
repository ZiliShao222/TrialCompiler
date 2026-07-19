"""Build the submission-ready TrialCompiler bibliography in Word format."""

from __future__ import annotations

import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CATALOG = ROOT / "references" / "catalog" / "sources.jsonl"
OUTPUT = ROOT / "docs" / "TrialCompiler_项目完整参考文献.docx"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 96, 105)

COLLECTIONS = {
    "01_regulatory_and_ethics": "法规、监管指南与研究伦理",
    "02_templates_and_associated_documents": "临床文件模板与关联文件",
    "03_data_standards_and_structured_protocols": "数据标准与结构化方案",
    "04_ai_methods_and_governance": "人工智能方法、不确定性、可解释性与治理",
    "05_industry_competitors_and_value": "行业产品、竞品与价值证据",
    "06_platform_and_integration": "平台与系统集成",
    "08_internal_project_materials": "项目内部材料（不作为外部证据）",
}


def style_run(run, size: float, *, bold: bool = False, color=None) -> None:
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(paragraph.add_run("第 "), 9, color=GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    style_run(paragraph.add_run(" 页"), 9, color=GRAY)


def add_hyperlink(paragraph, label: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([props, text])
    link.append(run)
    paragraph._p.append(link)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Title", 25, DARK_BLUE, 0, 8),
        ("Subtitle", 12, GRAY, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def main() -> int:
    sources = [json.loads(line) for line in CATALOG.read_text(encoding="utf-8").splitlines() if line.strip()]
    grouped: dict[str, list[dict[str, object]]] = defaultdict(list)
    for source in sources:
        grouped[str(source["collection"])].append(source)
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin, section.bottom_margin = Inches(0.85), Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.9)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(header.add_run("TRIALCOMPILER  ·  REFERENCES"), 8.5, bold=True, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("TrialCompiler 项目完整参考文献")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("法规 · 临床文件 · 数据标准 · AI方法 · 不确定性与可解释性 · 行业证据")
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(summary.add_run(f"规范化来源 {len(sources)} 条  |  按主题分类并保留原始链接"), 10, bold=True, color=DARK_BLUE)

    doc.add_heading("引用与证据边界", level=1)
    for principle in (
        "法规、伦理、统计与文件要求优先引用监管机构、国际标准组织及正式指南的原始页面。",
        "人工智能、不确定性与可解释性方法引用原始论文、会议或正式出版页面。",
        "厂商材料只支持功能和市场比较，不支持监管合规或临床有效性结论。",
        "ClinicalTrials.gov登记和公开附件在本项目评测中作为冻结正确基线；受控变异副本才是缺陷正例。",
        "INTERNAL材料只记录项目形成过程，不作为外部事实依据。",
    ):
        add_bullet(doc, principle)

    number = 0
    for collection, heading in COLLECTIONS.items():
        doc.add_heading(heading, level=1)
        for source in sorted(grouped.get(collection, []), key=lambda item: str(item["source_id"])):
            number += 1
            citation = doc.add_paragraph()
            citation.paragraph_format.left_indent = Inches(0.28)
            citation.paragraph_format.first_line_indent = Inches(-0.28)
            citation.paragraph_format.space_after = Pt(2)
            citation.paragraph_format.keep_together = True
            purpose = str(source.get("intended_use", "")).strip()
            citation.paragraph_format.keep_with_next = bool(purpose)
            style_run(citation.add_run(f"{number}. [{source['source_id']}] {source['title']}. "), 10.5, bold=True)
            style_run(citation.add_run(f"{source['organization']}, {source['year']}。"), 10.5)
            url = str(source.get("source_url", "")).strip()
            if url:
                citation.add_run(" ")
                add_hyperlink(citation, "原始来源", url)
            if purpose:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.28)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_together = True
                style_run(p.add_run("项目用途："), 9.5, bold=True, color=DARK_BLUE)
                style_run(p.add_run(purpose), 9.5, color=GRAY)

    doc.add_heading("数据来源与评测基线说明", level=1)
    doc.add_paragraph(
        "本项目真实案例来自ClinicalTrials.gov API v2、研究登记页及公开附件。未修改的登记快照、研究方案、"
        "统计分析计划和知情同意文件在评测中统一视为正确基线；系统只在隔离副本中应用有记录、可逆的"
        "受控变异，构造答案确定的缺陷正例。公开原始材料中的关键词和表面差异作为复杂负对照，用于检验误报控制。"
    )
    doc.add_heading("可复现性", level=1)
    doc.add_paragraph(
        "机器可读总账位于 references/catalog/sources.jsonl；来源别名、重复项、本地快照和校验信息位于 "
        "references/catalog/ 与 references/metadata/。本Word文件由脚本自动生成。"
    )
    doc.core_properties.title = "TrialCompiler 项目完整参考文献"
    doc.core_properties.subject = "项目提交统一参考文献表"
    doc.core_properties.author = "TrialCompiler 项目组"
    doc.save(OUTPUT)
    print(json.dumps({"output": str(OUTPUT), "source_count": len(sources)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
