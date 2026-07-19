"""Update the engineering Word report with corrected round-3 evidence."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "TrialCompiler_工科报告_6.4至结尾_Word版.docx"
OUTPUT = ROOT / "TrialCompiler_工科报告_6.4至结尾_正式稿.docx"


def find_heading(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Heading not found: {prefix}")


def delete_between(start: Paragraph, end: Paragraph) -> None:
    node = start._p.getnext()
    while node is not None and node is not end._p:
        following = node.getnext()
        node.getparent().remove(node)
        node = following


def insert_body(anchor: Paragraph, paragraphs: list[str]) -> None:
    for text in paragraphs:
        paragraph = anchor.insert_paragraph_before(text)
        paragraph.style = "Normal"


def delete_from(start: Paragraph) -> None:
    body = start._p.getparent()
    node = start._p
    while node is not None:
        following = node.getnext()
        if not node.tag.endswith("sectPr"):
            body.remove(node)
        node = following


def main() -> int:
    doc = Document(TARGET)

    heading_64 = find_heading(doc, "6.4 ")
    next_64 = Paragraph(heading_64._p.getnext(), heading_64._parent)
    direction_note = (
        "在本报告的受控评测中，当前案例的 ClinicalTrials.gov 冻结登记值始终作为权威事实；"
        "从其他案例移植的同类字段值只存在于隔离的缺陷副本中。语义模型不接收金标准标签，"
        "而是从包含多个字段的来源证据片段中定位目标事实、比较候选文档并形成可复核的修订建议。"
    )
    if direction_note not in [p.text for p in doc.paragraphs]:
        inserted = next_64.insert_paragraph_before(direction_note)
        inserted.style = "Normal"

    heading_71 = find_heading(doc, "7.1 ")
    heading_72 = find_heading(doc, "7.2 ")
    delete_between(heading_71, heading_72)
    insert_body(
        heading_72,
        [
            "仓库已经实现事实和来源契约、依赖图、多角色审核、分阶段方案生成、确定性检查、Qwen 语义复核、原子修订、沙箱复检、专业决策请求、审计记录、独立评分程序、命令行工具和应用程序编程接口。确定性路径处理数值、枚举值、标识符和旧值残留等可形式化问题；语义路径从冻结证据片段中提取目标事实并形成带理由的候选判断。两类结果进入同一修订和复检流程，但语言模型不能修改权威事实、批准状态或发布门禁。",
            "验证材料只使用公开研究层资料或完全合成材料，不包含患者级数据。这里的“完成”表示研究原型与可复现实验链已经打通，不表示企业生产部署、法规认可或临床有效性验证已经完成。相关回归测试已全部通过；完整测试数量以最终冻结提交的持续集成记录为准。",
        ],
    )

    heading_73 = find_heading(doc, "7.3 ")
    delete_between(heading_72, heading_73)
    insert_body(
        heading_73,
        [
            "原型验证围绕四个递进问题展开：输入资料能否真实、可读并冻结复现；确定性检查能否发现预定义字段缺陷并保护未修改负对照；语义模型能否从带干扰字段的登记证据中定位目标事实、判断候选差异并提出修订；获批候选能否完成局部修改、保持来源链且不引入新问题。四类问题分别对应输入可靠性、规则检测能力、证据提取与语义一致性能力以及修订安全性，结论不得跨层外推。",
            "评测遵循来源和评分规则预先确定、案例级隔离、正负对照并存以及生成与评分分离的原则。50 个案例按研究编号划分为 30 个开发案例、10 个校准案例和 10 个独立测试案例。同一研究的注册快照、公开附件、候选值、负对照和修订记录始终位于同一集合，避免同一研究的信息跨集合泄漏。开发集用于规则编写和错误分析，校准集用于冻结模型输入格式和评分规则，测试集在策略冻结后用于独立评价。",
            "输入可靠性由 50 份注册快照、65 份官方 PDF、文件摘要、页数和文本可提取性验证。任务有效性分别由官方文档元数据、400 个跨试验同字段错配缺陷与 400 个未修改负对照、以及真实 Qwen API 盲测评价。修订安全性通过 400 个受控缺陷的原子补丁和沙箱复检评价，检查权威值写入、错误值移除、修改局部性、来源保持、问题关闭和新增问题。",
            "确定性检查结果、Qwen 语义复核结果和自然语言风险候选不合并为单一“系统准确率”。确定性实验衡量已编码字段约束，Qwen 实验衡量结构化证据提取与一致性判断；自然临床语义仍需独立专业裁决。三者的输入难度、金标准形成方式和适用边界不同，必须分别报告。",
            "对二元结果同时报告样本比例和 Wilson 得分法 95% 置信区间。置信区间下限是有限样本总体比例的保守端点，不是样本中的最低观测值；即使样本结果为 100%，也不能据此宣称总体零风险。",
        ],
    )

    heading_75 = find_heading(doc, "7.5 ")
    delete_from(heading_75)

    doc.add_heading("7.5 50 例公开案例上的受控缺陷检测、语义复核与修订闭环", level=2)
    for text in [
        "ClinicalTrials.gov 是由美国国家医学图书馆运营的公开临床研究注册与结果信息平台。其第二版应用程序编程接口以结构化形式提供研究编号、研究类型、招募状态、设计字段、结果信息和公开文档元数据。项目通过该接口冻结 50 个真实研究的注册记录，并下载注册页面链接的试验方案、统计分析计划等文件，形成 50 份注册快照和 65 份公开 PDF，共 3,025 页。50 个案例中，48 个为干预性研究、2 个为观察性研究；35 个案例使用合并的试验方案与统计分析计划，15 个使用分立文件。构建过程检查文档元数据、PDF 文件头、页数、正文可提取性和 SHA-256 内容摘要。",
        "ClinicalTrials.gov 登记快照和公开附件在受控一致性评测中作为冻结来源基线；这一评测约定用于确定字段传播问题的答案，不等同于美国国家医学图书馆、美国国立卫生研究院或监管机构对文件全部临床内容作出的专业批准。",
        "受控缺陷覆盖入组人数、正式研究标题、研究总体状态、研究类型、主要结局指标、主要结局时间范围、疾病或研究条件以及研究臂数量等 8 类字段。每个案例的原始登记值始终作为权威事实；缺陷候选通过将另一项真实试验的同类字段值移植到当前案例的隔离文档副本中构造，并配置一条保持当前登记值的未修改负对照。由此形成 400 个受控缺陷和 400 个负对照，共 800 个评价场景。每条记录均保存当前案例编号、权威值、错误候选值、错误值来源案例、注册快照摘要和数据集划分。该设计不把人工变异值误作权威事实，也不把受控构造缺陷宣称为公开数据库中自然存在的错误。",
        "TrialCompiler 对 800 个场景执行确定性字段一致性检查，得到真阳性 400、假阳性 0、假阴性 0、真阴性 400。400 个受控缺陷全部被检出，400 个负对照均未报警；全量准确率、精确率、召回率和 F1 值均为 100%。400 个缺陷检出率的 Wilson 得分法 95% 置信区间为 99.05%—100%，800 个场景总体准确率的 95% 置信区间为 99.52%—100%。独立测试集包含 80 个缺陷和 80 个负对照，共 160 个场景，均被正确分类；总体准确率的 95% 置信区间为 97.66%—100%。这些数字验证的是 8 类预定义字段在冻结来源和确定判断边界内的一致性检查能力，不表示自然临床语义审核总体准确率为 100%。",
        "对 400 个受控缺陷，系统依据冻结标签触发测试性批准，生成原子修订并执行沙箱复检。闭环成功要求权威登记值写入目标单元、错误候选值从目标范围移除、修改不超出批准单元、来源关系保持完整、原问题关闭且不产生新问题。400 个缺陷均完成正确修订和复检，修订闭环成功率为 100%，95% 置信区间为 99.05%—100%；400 个负对照均保持不变。这里的“批准”是基于受控金标准触发测试流程，不是医学、统计或注册专家对自然临床问题作出的正式签字。",
        "为评价语义模型能否从证据中独立定位字段，而非只执行代码级字符串规则，项目通过阿里云百炼 DashScope 的 OpenAI 兼容接口调用 qwen-plus 开展盲测。模型输入包含当前试验编号、目标字段名称、含 8 类字段的冻结注册证据片段和候选文档陈述，但不包含缺陷标签、缺陷来源案例或预设替换答案。模型需要自行定位目标字段、提取权威值、判断候选内容是否一致，并在需要时给出修订值。",
        "校准集使用 16 个场景覆盖 8 类字段的缺陷和负对照，用于验证接口、输出结构和数值单位归一化规则；随后冻结系统提示词、输入结构、数据标签和评分规则。独立测试集执行 160 次真实 API 请求，全部成功并产生 160 个唯一请求编号，实际墙钟时间为 181.81 秒，平均单请求延迟为 4.50 秒，共消耗 83,405 个 token。",
        "在 80 个受控缺陷和 80 个负对照中，Qwen 得到真阳性 80、假阳性 0、假阴性 0、真阴性 80，分类准确率、精确率、召回率和 F1 值均为 100%，总体准确率的 Wilson 得分法 95% 置信区间为 97.66%—100%。模型在 159/160 个场景中给出了与冻结登记逐字一致的替换值，精确修订率为 99.38%。",
        "唯一精确修订差异出现在主要结局时间范围字段。冻结登记原文将 approximately 拼写为 approximatly；模型正确识别候选时间范围与权威证据冲突，却在建议修订时自动纠正了来源中的拼写。该结果在语义上保持正确，但不满足逐字保留冻结来源的严格追溯标准，因此计为修订值错误。测试结束后未调整提示词或重跑以消除该错误。",
        "上述 Qwen 结果证明模型能够从结构化登记证据片段中提取指定字段，并识别受控的跨试验同字段错配；它不等同于对完整试验方案、统计分析计划、知情同意书或自然发生临床缺陷的专家级总体准确率。语言模型输出仍属于候选审核意见，不能替代医学、统计、注册或质量人员的正式决定。",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("7.6 GitHub、Demo 与复现方式", level=2)
    for text in [
        "代码仓库地址为 https://github.com/ZiliShao222/TrialCompiler，本文实验结果对应提交 cc69d9a。复现应优先使用仓库内冻结输入和逐条记录，而不是先联网重新下载持续变化的注册数据。",
        "确定性评测可执行：python scripts/run_round3_rich_defect_benchmark.py。输出位于 benchmarks/trialdocbench/public_corpus_050/round3_rich_defects_v2/，包含 800 条逐项记录、汇总报告和设计说明。",
        "真实 Qwen 复现需从本地安全环境向 DASHSCOPE_API_KEY 注入密钥，先以 --split calibration --limit 16 --workers 4 运行校准，再以 --split test --workers 4 运行测试。密钥不写入仓库、配置文件、运行记录或审计日志。冻结输出位于 benchmarks/trialdocbench/public_corpus_050/round3_qwen_api_v2/，保存逐条模型输出、请求编号、延迟、token 用量和汇总指标。",
        "由于 qwen-plus 是服务端模型，冻结 JSON 可以逐项复核，但未来在同名模型上重新请求不保证回答文字逐字相同。ClinicalTrials.gov 在线重建也应创建新语料版本，不得静默覆盖已经标注的冻结 corpus。",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("7.7 后续专业验证与扩围条件", level=2)
    for text in [
        "当前受控实验已经完成字段一致性检测、Qwen 证据复核和原子修订闭环验证。后续专业验证不作为当前受控指标成立的前置条件，而用于评价系统在自然发生的复杂临床语义、跨文档逻辑、统计策略和注册判断中的适用范围。",
        "项目另保留了带来源定位的自然语言风险候选，用于后续分层抽样和专业验证。在医学、统计、注册或质量人员形成独立金标准前，这些候选只作为风险定位和人工路由能力展示，不与本报告的受控缺陷结果合并，也不报告准确率。",
        "后续验证将扩展不同治疗领域、语言、文档类型和企业模板，并比较纯人工流程与 TrialCompiler 辅助流程的问题发现、错误修订、影响范围、审核时间和返工情况。预测阈值和处置策略只能在校准集选择，独立测试标签在策略冻结后使用；若服务端模型、提示词、来源快照或评分规则发生变化，应建立新的版本化实验，不得覆盖既有结果。",
        "进入企业试点前仍需完成权限隔离、安全测试、企业模板与标准操作程序适配、正式专业复核和质量体系验收。真实患者数据、未经授权的企业资料、自动批准研究设计以及自动发布正式临床文件均不属于当前原型的能力范围。",
    ]:
        doc.add_paragraph(text)

    doc.save(OUTPUT)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
